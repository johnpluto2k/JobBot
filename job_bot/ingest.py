"""Document ingestion: walk the docs folder and extract raw text.

Supports PDF (pymupdf4llm -> markdown, fallback to plain PyMuPDF), DOCX
(python-docx), and plain text/markdown. Returns a list of ParsedDoc objects
that the extraction layer turns into a structured profile.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from . import config


@dataclass
class ParsedDoc:
    path: Path
    text: str
    parser: str
    doc_type: str = "other"

    @property
    def filename(self) -> str:
        return self.path.name

    @property
    def char_count(self) -> int:
        return len(self.text)


def _classify(path: Path) -> str:
    """Rough document-type guess from the path/filename."""
    p = str(path).lower()
    name = path.name.lower()
    # Filename-based checks win over folder-based ones (a project file can live
    # in a "Resumes" folder).
    if "cover" in name and "letter" in name:
        return "cover_letter"
    if "transcript" in name:
        return "transcript"
    if "project" in name:
        return "project"
    if "coverletter" in p or "cover letter" in p:
        return "cover_letter"
    if "resume" in p or "cv" in name:
        return "resume"
    if "transcript" in p:
        return "transcript"
    return "other"


def _read_pdf(path: Path) -> tuple[str, str]:
    try:
        import pymupdf4llm

        return pymupdf4llm.to_markdown(str(path)), "pymupdf4llm"
    except Exception:
        pass
    # Fallback: plain text extraction via PyMuPDF.
    import fitz  # PyMuPDF

    doc = fitz.open(str(path))
    try:
        return "\n".join(page.get_text() for page in doc), "pymupdf"
    finally:
        doc.close()


def _read_docx(path: Path) -> tuple[str, str]:
    import docx  # python-docx

    document = docx.Document(str(path))
    parts: list[str] = [p.text for p in document.paragraphs]
    # Tables hold real content in some resumes/cover letters.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts), "python-docx"


def _read_text(path: Path) -> tuple[str, str]:
    return path.read_text(encoding="utf-8", errors="replace"), "text"


def _clean(text: str) -> str:
    # Collapse runs of blank lines but keep paragraph structure.
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_file(path: Path) -> ParsedDoc | None:
    """Parse a single supported file into a ParsedDoc, or None if unsupported/empty."""
    suffix = path.suffix.lower()
    try:
        if suffix in config.PDF_SUFFIXES:
            text, parser = _read_pdf(path)
        elif suffix in config.DOCX_SUFFIXES:
            text, parser = _read_docx(path)
        elif suffix in config.TEXT_SUFFIXES:
            text, parser = _read_text(path)
        else:
            return None
    except Exception as exc:  # keep one bad file from killing the run
        print(f"  ! failed to parse {path.name}: {exc}")
        return None

    text = _clean(text)
    if not text:
        return None
    return ParsedDoc(path=path, text=text, parser=parser, doc_type=_classify(path))


def discover_files(docs_dir: Path | None = None) -> list[Path]:
    """Find all supported documents under docs_dir (recursive)."""
    base = Path(docs_dir) if docs_dir else config.DOCS_DIR
    if not base.exists():
        return []
    files: list[Path] = []
    for path in sorted(base.rglob("*")):
        if path.is_file() and path.suffix.lower() in config.SUPPORTED_SUFFIXES:
            # Skip temporary Office lock files.
            if path.name.startswith("~$"):
                continue
            files.append(path)
    return files


def ingest(docs_dir: Path | None = None, save_raw: bool = True) -> list[ParsedDoc]:
    """Parse every supported document under docs_dir."""
    files = discover_files(docs_dir)
    parsed: list[ParsedDoc] = []
    for path in files:
        doc = parse_file(path)
        if doc is None:
            continue
        parsed.append(doc)
        print(f"  + {doc.filename}  [{doc.doc_type}, {doc.char_count} chars via {doc.parser}]")
        if save_raw:
            config.ensure_dirs()
            out = config.RAW_TEXT_DIR / (path.stem + ".txt")
            out.write_text(doc.text, encoding="utf-8")
    return parsed
