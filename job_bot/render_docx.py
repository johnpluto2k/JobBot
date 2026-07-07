"""Render a TailoredResume to an ATS-clean one-page DOCX.

Format rules (from the master plan): single column, no tables/text boxes/
graphics, standard font (Calibri 11), black only, 0.5"+ margins, strong action
verbs, no pronouns.
"""

from __future__ import annotations

from pathlib import Path

from .resume_models import TailoredResume
from .writing_style import split_metrics

FONT = "Calibri"
BODY_PT = 10.5
NAME_PT = 18
HEAD_PT = 11.5


def render_docx(resume: TailoredResume, out_path: Path) -> Path:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(36)      # 0.5"
        section.left_margin = section.right_margin = Pt(54)      # 0.75"

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(BODY_PT)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    def _spacing(p, before=0.0, after=0.0):
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)

    # --- Header ---
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run((resume.name or "").upper())
    run.bold = True
    run.font.size = Pt(NAME_PT)
    _spacing(h, after=2)

    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = c.add_run(resume.contact_line)
    cr.font.size = Pt(BODY_PT)
    _spacing(c, after=6)

    def section_heading(text: str) -> None:
        p = doc.add_paragraph()
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(HEAD_PT)
        _spacing(p, before=6, after=2)
        # bottom border line
        pPr = p._p.get_or_add_pPr()
        from docx.oxml import OxmlElement
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pbdr.append(bottom)
        pPr.append(pbdr)

    def entry_header(left: str, right: str | None, italic_left: str | None = None) -> None:
        p = doc.add_paragraph()
        # tab stop at right margin for the date
        from docx.enum.text import WD_TAB_ALIGNMENT
        from docx.shared import Inches
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
        rl = p.add_run(left)
        rl.bold = True
        if right:
            p.add_run("\t" + right).bold = True
        _spacing(p, before=4, after=0)
        if italic_left:
            sp = doc.add_paragraph()
            si = sp.add_run(italic_left)
            si.italic = True
            _spacing(sp, after=1)

    def bullet(text: str) -> None:
        p = doc.add_paragraph(style="List Bullet")
        # Playbook: bold all metrics — recruiters scan numbers first.
        for chunk, is_metric in split_metrics(text):
            p.add_run(chunk).bold = is_metric
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Pt(12)

    if resume.summary:
        section_heading("Summary")
        s = doc.add_paragraph()
        s.add_run(resume.summary)
        _spacing(s, after=2)

    # --- Education ---
    if resume.education:
        section_heading("Education")
        for e in resume.education:
            right = e.graduation_date
            entry_header(e.school or "", right)
            deg = doc.add_paragraph()
            line = e.degree or ""
            if e.gpa:
                line += f"   |   GPA: {e.gpa}"
            deg.add_run(line)
            _spacing(deg, after=0)
            if e.secondary:
                sp = doc.add_paragraph(); sp.add_run(e.secondary); _spacing(sp, after=0)
            if e.coursework:
                cw = doc.add_paragraph()
                cw.add_run("Relevant Coursework: ").bold = True
                cw.add_run(", ".join(e.coursework))
                _spacing(cw, after=2)

    def role_section(title: str, roles) -> None:
        if not roles:
            return
        section_heading(title)
        for role in roles:
            left = role.organization or role.role or ""
            sub = None
            if role.role and role.organization:
                sub = role.role + (f"  |  {role.location}" if role.location else "")
            entry_header(left, role.dates, italic_left=sub)
            for b in role.bullets:
                bullet(b)

    role_section("Experience", resume.experience)
    role_section("Leadership & Activities", resume.leadership)

    if resume.projects:
        section_heading("Projects")
        for p in resume.projects:
            entry_header(p.organization or p.role or "", None)
            for b in p.bullets:
                bullet(b)

    if resume.skills:
        section_heading("Skills")
        sp = doc.add_paragraph()
        sp.add_run(", ".join(resume.skills))
        _spacing(sp, after=0)
    if resume.certifications:
        cp = doc.add_paragraph()
        cp.add_run("Certifications: ").bold = True
        cp.add_run(", ".join(resume.certifications))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
